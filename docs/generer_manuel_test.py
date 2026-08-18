"""
Génère le manuel de test / recette Educ_RDC (Word, checklists par rôle).
Usage : python docs/generer_manuel_test.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLEU_NUIT = RGBColor(0x0B, 0x2A, 0x4A)
BLEU = RGBColor(0x1A, 0x56, 0x8C)
GRIS = RGBColor(0x5A, 0x67, 0x72)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
VERT = RGBColor(0x1F, 0x7A, 0x4D)
ROUGE = RGBColor(0xC4, 0x1E, 0x3A)

VERSION = '1.0'
DATE_DOC = date.today().strftime('%d/%m/%Y')
CASES = '☐  OK     ☐  NOK     ☐  N/A'


def set_run_font(run, size=11, bold=False, color=None, name='Calibri'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn('w:shd'):
            tcPr.remove(child)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, color='1A568C', size='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=GRIS)


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size, color in (
        (1, 16, BLEU_NUIT),
        (2, 13, BLEU),
        (3, 11.5, BLEU_NUIT),
    ):
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(
        f'Educ_RDC  ·  Manuel de test v{VERSION}  ·  {DATE_DOC}  ·  Usage interne'
    )
    set_run_font(run, size=9, color=GRIS)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run('MINEDUC-NC  —  Recette applicative  —  Page ')
    set_run_font(r1, size=9, color=GRIS)
    add_page_number(fp)


def add_horizontal_line(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A568C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_para(doc, text, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def fill_cell(cell, text, *, bold=False, size=9.5, color=None, fill=None, center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold, color=color)
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell, 'C5D0DC', '4')
    set_cell_margins(cell)


def add_checklist(doc, items, start=1):
    """items : list[str] ou list[tuple(controle, resultat_attendu)]"""
    table = doc.add_table(rows=1 + len(items), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ('N°', 'Contrôle à réaliser', 'Résultat', 'Observations')
    widths = (Cm(1.1), Cm(9.4), Cm(3.6), Cm(3.2))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        fill_cell(cell, h, bold=True, size=9, color=BLANC, fill='0B2A4A', center=True)
        cell.width = widths[i]
    for idx, item in enumerate(items):
        if isinstance(item, (list, tuple)) and len(item) >= 2:
            controle, attendu = item[0], item[1]
            texte = f'{controle}\nAttendu : {attendu}'
        else:
            texte = str(item)
        row = table.rows[idx + 1]
        alt = 'F4F7FA' if idx % 2 else 'FFFFFF'
        fill_cell(row.cells[0], f'{start + idx:02d}', bold=True, size=9, fill=alt, center=True)
        fill_cell(row.cells[1], texte, size=9, fill=alt)
        fill_cell(row.cells[2], CASES, size=8, color=GRIS, fill=alt, center=True)
        fill_cell(row.cells[3], '', size=9, fill=alt)
        for i, w in enumerate(widths):
            row.cells[i].width = w
    doc.add_paragraph()
    return start + len(items)


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, bold=True, size=9.5, color=BLANC, fill='0B2A4A', center=True)
    for r_idx, row in enumerate(rows):
        fill = 'F4F7FA' if r_idx % 2 else 'FFFFFF'
        for c_idx, value in enumerate(row):
            fill_cell(table.rows[r_idx + 1].cells[c_idx], value, size=9.5, fill=fill)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    r = p.add_run(f'▶ {title}  ')
    set_run_font(r, size=10, bold=True, color=BLEU)
    r2 = p.add_run(text)
    set_run_font(r2, size=10, color=GRIS)


def add_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('RÉPUBLIQUE DÉMOCRATIQUE DU CONGO')
    set_run_font(r, size=12, bold=True, color=BLEU_NUIT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Ministère de l’Éducation Nationale et de Nouvelle Citoyenneté')
    set_run_font(r, size=11, color=BLEU)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run('MANUEL DE TEST')
    set_run_font(r, size=28, bold=True, color=BLEU_NUIT, name='Calibri Light')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Recette applicative — checklists par rôle')
    set_run_font(r, size=16, color=BLEU)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Plateforme Educ_RDC')
    set_run_font(r, size=14, bold=True, color=BLEU_NUIT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run(
        f'Document de recette  ·  Version {VERSION}  ·  {DATE_DOC}\n'
        'Destination : testeurs, maîtrise d’ouvrage, exploitation'
    )
    set_run_font(r, size=11, color=GRIS)

    for _ in range(3):
        doc.add_paragraph()

    make_table(
        doc,
        ['Élément', 'Valeur'],
        [
            ['Application', 'Educ_RDC — identification scolaire nationale'],
            ['Périmètre', 'Tous les rôles et modules en production'],
            ['Type de tests', 'Recette fonctionnelle (UI) + contrôles d’accès'],
            ['Environnement', 'À renseigner (recette / préprod / local)'],
            ['URL testée', ''],
            ['Jeu de données', 'Année scolaire active + école / classe témoins'],
        ],
        col_widths=[5.0, 12.3],
    )
    doc.add_page_break()


def add_toc(doc: Document):
    doc.add_heading('Sommaire', level=1)
    items = [
        '1. Objet, conventions et prérequis',
        '2. Synthèse des rôles',
        '3. Tests transversaux (tous rôles)',
        '4. Administrateur',
        '5. Agent National',
        '6. Agent Province administrative',
        '7. Agent Province éducationnelle',
        '8. Agent Antenne',
        '9. Administratif école',
        '10. Enseignant',
        '11. Tests d’interdiction et non-régression',
        '12. Fiche de recette et visas',
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=12, color=BLEU_NUIT)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Contenu des checklists
# ---------------------------------------------------------------------------

TRANSVERSAL = [
    (
        'Connexion — identifiants valides',
        'Accès au tableau de bord ; message de bienvenue ; menu adapté au rôle.',
    ),
    (
        'Connexion — identifiants invalides',
        'Message « Identifiants incorrects » ; aucun accès.',
    ),
    (
        'Session unique — 2e connexion pendant qu’une session est en ligne',
        'Refus : « Ce compte a déjà une session en ligne… » ; la 1re session reste ouverte.',
    ),
    (
        'Session unique — reconnexion après déconnexion',
        'La nouvelle connexion est acceptée.',
    ),
    (
        'Un seul onglet actif — ouvrir un 2e onglet Educ_RDC',
        'L’ancien onglet affiche « Page inactive » ; seul le nouvel onglet fonctionne.',
    ),
    (
        'Réactiver une page inactive',
        'Le bouton « Activer cette page » bascule l’activité ; l’autre onglet se met en pause.',
    ),
    (
        'PDF (fiche élève, liste, bulletin) en nouvel onglet',
        'Le PDF s’ouvre sans bloquer l’application (exception volontaire).',
    ),
    (
        'Mon profil — consulter / modifier e-mail et téléphone',
        'Enregistrement OK ; avatar et nom affichés dans le bandeau.',
    ),
    (
        'Changement de mot de passe depuis Mon profil',
        'Succès ; reconnexion possible avec le nouveau mot de passe.',
    ),
    (
        'Déconnexion',
        'Retour à l’écran de connexion ; l’URL /dashboard/ redirige vers login.',
    ),
    (
        'Biométrie (si le compte a l’option activée) — enregistrement dans Mon profil',
        'Appareil enregistré (HTTPS ou localhost).',
    ),
    (
        'Biométrie — connexion depuis l’écran login (identifiant + empreinte / Face ID)',
        'Accès dashboard. Sans autorisation biométrique : refus.',
    ),
    (
        'Accès hors RDC (si IP étrangère et non admin)',
        'Blocage + demande enregistrée pour l’administrateur.',
    ),
]


ADMIN = {
    'intro': (
        'Périmètre national. Accès complet hors module Évaluation (réservé à l’école). '
        'Compte de test : rôle Administrateur.'
    ),
    'menus': [
        (
            'Menus visibles',
            'Tableau de bord, Écoles, Élèves, Gestion documentaire, Rapports, '
            'Gestion Utilisateurs, Utilisateurs connectés, Paramètres. Pas de menu Évaluation.',
        ),
        (
            'URL /evaluations/ saisie manuellement',
            'Pas d’accès métier évaluation (redirection ou refus).',
        ),
    ],
    'dashboard': [
        ('Tableau de bord national', 'KPI écoles / élèves ; graphiques ; listes d’effectifs.'),
    ],
    'ecoles': [
        ('Liste des écoles — recherche et filtres (province, antenne, type, niveau)', 'Résultats cohérents.'),
        ('Créer une école (identité, code, type, niveau, PE/antenne, contacts, GPS)', 'Fiche créée ; visible en liste.'),
        ('Modifier / désactiver une école', 'Modifications persistées.'),
        ('Fiche école — photos, documents, personnel, import personnel', 'Ajout / suppression OK.'),
        ('Import Excel d’écoles (modèle national)', 'Rapport d’import ; pas de doublon de code.'),
        ('Structure scolaire — affecter section / option / classes EPSP à une école', 'Classes créées (ex. Latin-Philosophie 1ère humanités).'),
    ],
    'eleves': [
        ('Créer un élève (identité, parents, classe, photo)', 'Fiche complète ; QR généré une seule fois.'),
        ('Modifier un élève hors QR', 'QR inchangé (immuable).'),
        ('Import Excel d’élèves', 'Affectation classe / école correcte.'),
        ('Fiche élève — PDF, QR, cartes, biométrie', 'Blocs visibles ; PDF s’ouvre.'),
        ('Suppression / désactivation élève', 'Plus listé parmi les actifs.'),
    ],
    'users': [
        ('Créer un agent (PA / PE / antenne / national) avec rattachement', 'Compte utilisable ; périmètre respecté.'),
        ('Créer un enseignant avec classe titulaire', 'Classe déjà titulaire : refus « a déjà un titulaire ».'),
        ('Classes occupées grisées dans le formulaire', 'Impossible de sélectionner une classe déjà attribuée.'),
        ('Activer / désactiver un compte et l’option biométrique', 'Connexion refusée si inactif ; biométrie selon le flag.'),
        ('Page Gestion de permissions + export PDF', 'Matrice lisible, conforme aux droits réels.'),
    ],
    'monitoring': [
        ('Utilisateurs connectés — sessions / uniques / en ligne (15 min)', 'La répartition compte les utilisateurs uniques, pas les cookies en trop.'),
        ('Filtre En ligne vs Session active', 'Le tableau suit le filtre ; une session inactive n’est pas « en ligne ».'),
        ('Historique d’un utilisateur (bouton Historique)', 'Journal : connexions, déconnexions, refus, IP ; filtre et pagination.'),
        ('Forcer la déconnexion d’une session (pas la sienne)', 'L’utilisateur cible est renvoyé au login.'),
        ('Carte des sessions localisées en RDC', 'Points GPS/IP ; hors RDC masqué.'),
        ('Accès hors RDC — autoriser 7 j / toutes IP / refuser / révoquer', 'Décision persistée ; connexion étrangère ensuite conforme.'),
    ],
    'params': [
        ('Paramètres — PA / PE / Antennes (création, organigramme)', 'Hiérarchie cohérente.'),
        ('Années scolaires — une seule année active ; périodes créées', 'Régime primaire vs secondaire respecté.'),
        ('Catalogue matières (Structure scolaire) pour une option / classe', 'Matières rattachées à la classe (pas Coupe et Couture sur Latin-Philosophie).'),
        ('Gestion documentaire — arrêtés / documents nationaux', 'Dépôt et consultation OK.'),
        ('Rapports nationaux + exports PDF', 'Données nationales ; pas d’erreur.'),
    ],
}


AGENT_NAT = {
    'intro': (
        'Même périmètre métier que l’administrateur (écoles, élèves, référentiels), '
        'sans gestion des comptes, sans monitoring, sans évaluation.'
    ),
    'items': [
        ('Menus', 'Écoles, Élèves, Gestion documentaire, Rapports, Paramètres. Pas Utilisateurs, pas Monitoring, pas Évaluation.'),
        ('Créer / modifier école et élève', 'Écriture autorisée.'),
        ('Import Excel écoles / élèves', 'Autorisé.'),
        ('Catalogue matières via Structure scolaire', 'Chargement du catalogue d’option (ex. LP).'),
        ('URL /utilisateurs/ ou /monitoring/utilisateurs-connectes/', 'Accès refusé ou redirection.'),
        ('Création de compte utilisateur', 'Impossible.'),
    ],
}


AGENT_PA = {
    'intro': (
        'Chef des PE de sa province administrative. Consultation territoriale uniquement. '
        'Le compte doit être rattaché à une PA.'
    ),
    'items': [
        ('Dashboard', 'KPI des PE / écoles / élèves de la PA uniquement.'),
        ('Liste écoles — filtres', 'Filtres PE, antenne, type, niveau. Aucune école hors PA.'),
        ('Fiche école', 'Lecture (photos, docs, personnel) ; pas de bouton Créer / Modifier / Supprimer.'),
        ('Élèves', 'Lecture dans le périmètre PA ; pas de création / import / suppression.'),
        ('Photo élève', 'Modification refusée.'),
        ('Rapports', 'Vue PA ; pas de cartes / biométries en écriture.'),
        ('Menus absents', 'Évaluation, Paramètres, Utilisateurs, Monitoring, Gestion documentaire.'),
    ],
}


AGENT_PE = {
    'intro': 'Consultation de sa province éducationnelle. Compte rattaché à une PE.',
    'items': [
        ('Dashboard', 'KPI antennes / écoles / élèves de la PE.'),
        ('Filtres écoles', 'Antenne + type + niveau (pas de choix de PA). Hors PE : aucune école.'),
        ('Écoles / élèves', 'Lecture seule dans la PE.'),
        ('Écriture école / élève / notes', 'Refusée.'),
        ('Menus absents', 'Évaluation, Paramètres, Utilisateurs, Monitoring.'),
    ],
}


AGENT_ANTENNE = {
    'intro': 'Consultation de son antenne uniquement.',
    'items': [
        ('Dashboard', 'KPI écoles / élèves de l’antenne.'),
        ('Filtres écoles', 'Type + niveau seulement (pas de filtre antenne / PE).'),
        ('Périmètre', 'Aucune école d’une autre antenne.'),
        ('Écoles / élèves', 'Lecture seule.'),
        ('Menus absents', 'Évaluation, Paramètres, Utilisateurs, Monitoring, Gestion documentaire.'),
    ],
}


ADMIN_ECOLE = {
    'intro': (
        'Gère uniquement son établissement : fiche, élèves, programme d’évaluation, '
        'comptes admin_ecole / enseignant. Pas de saisie des notes.'
    ),
    'menus': [
        ('Menus', 'Tableau de bord, Mon école (pas la liste nationale), Élèves, Évaluation, Rapports. Pas Paramètres nationaux, pas Monitoring.'),
        ('URL /ecoles/ liste nationale', 'Redirection ou périmètre limité à son école.'),
    ],
    'ecole': [
        ('Mon école — modifier la fiche (sauf rattachement territorial national)', 'Enregistrement OK.'),
        ('Photos et documents d’école', 'Ajout / suppression OK.'),
        ('Personnel — CRUD + import', 'Limité à l’école.'),
        ('Comptes école — créer un administratif école', 'Rôle limité ; école imposée.'),
        ('Comptes — créer un enseignant titulaire', 'Une classe = un titulaire ; classes occupées désactivées.'),
        ('Créer un agent national / PE', 'Refusé.'),
    ],
    'eleves': [
        ('Créer / modifier / importer des élèves de l’école', 'Classe de l’école uniquement.'),
        ('Élève d’une autre école (URL directe)', 'Accès refusé.'),
        ('Photo élève', 'Modifiable.'),
        ('Blocs cartes / biométrie sur la fiche', 'Masqués.'),
    ],
    'eval': [
        ('Évaluation — choisir année + classe (ex. Littéraire · Latin-Philosophie — 1ère humanités)', 'Liste Matière (cours de la classe) se remplit (Latin, Philosophie, Français…).'),
        ('Aucune matière au programme — Appliquer le programme', 'Cours créés pour CETTE classe, pas ceux d’une autre classe / option.'),
        ('Consultation de la grille de notes', 'Lecture possible ; pas de saisie (champs inactifs ou absents).'),
        ('Clôturer / rouvrir une période pour la classe', 'Clôture = tous les cours ; réouverture admin_ecole OK.'),
        ('Classement et bulletins PDF', 'Génération OK pour la classe.'),
        ('Créer une matière (catalogue national)', 'Refusé — réservé à l’administration nationale.'),
    ],
    'rapports': [
        ('Rapports de l’école + PDF liste élèves / cours', 'Données de l’établissement uniquement.'),
    ],
}


ENSEIGNANT = {
    'intro': (
        'Titulaire d’une seule classe. Consulte ses élèves, met à jour leur photo, '
        'saisit les notes et imprime listes / bulletins de classe.'
    ),
    'items': [
        ('Menus', 'Tableau de bord, Élèves, Évaluation, Rapports (exports). Pas Écoles, pas Paramètres, pas Utilisateurs.'),
        ('Dashboard', 'Effectif de la classe (garçons / filles).'),
        ('Liste élèves', 'Uniquement la classe titulaire.'),
        ('Créer / importer / supprimer un élève', 'Refusé.'),
        ('Photo d’un élève de sa classe', 'Modifiable.'),
        ('Photo d’un élève hors classe (URL)', 'Refusé.'),
        ('Évaluation — matière', 'Uniquement les cours du programme de sa classe.'),
        ('Saisie des notes d’une période ouverte', 'Enregistrement ; max de la période respecté.'),
        ('Saisie sur période clôturée', 'Consultation seule ; enregistrement impossible.'),
        ('Clôturer la période (tous les cours de la classe)', 'Saisie bloquée pour toute la classe.'),
        ('Bulletin PDF / liste de cours PDF', 'S’ouvrent pour sa classe.'),
        ('Créer une matière ou charger le catalogue national', 'Refusé.'),
        ('Deux enseignants sur la même classe', 'Impossible (contrainte titulaire unique).'),
    ],
}


INTERDITS = [
    (
        'Enseignant — /ecoles/, /utilisateurs/, /monitoring/, /parametres/',
        'Pas d’accès (masqué + refus serveur).',
    ),
    (
        'Admin école — école d’un autre établissement (URL)',
        '404 / refus.',
    ),
    (
        'Agent antenne — école hors antenne',
        'Absente de la liste ; URL directe refusée.',
    ),
    (
        'Agent PE — filtre PA',
        'Filtre PA absent ; périmètre PE uniquement.',
    ),
    (
        'Agent national — monitoring et CRUD utilisateurs',
        'Refus.',
    ),
    (
        'Administrateur — saisie des notes',
        'Module évaluation non exposé.',
    ),
    (
        'Deux sessions en ligne pour le même login',
        'Toujours refusé tant que la 1re est en ligne (< 15 min d’activité).',
    ),
    (
        'Deux pages applicatives actives (onglets)',
        'Une seule page active ; l’autre en pause.',
    ),
]


def add_role_header(doc, numero, titre, intro, compte=''):
    doc.add_heading(f'{numero}. {titre}', level=1)
    add_para(doc, intro)
    make_table(
        doc,
        ['Champ', 'Valeur de test'],
        [
            ['Rôle', titre],
            ['Identifiant testé', compte or ''],
            ['Date / testeur', ''],
            ['Environnement / URL', ''],
            ['Verdict du chapitre', '☐  OK     ☐  NOK     ☐  Reporté'],
        ],
        col_widths=[5.0, 12.3],
    )


def build():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc(doc)

    # 1
    doc.add_heading('1. Objet, conventions et prérequis', level=1)
    add_para(
        doc,
        'Ce manuel sert à la recette fonctionnelle d’Educ_RDC. Chaque contrôle se '
        'coche ☐ OK (conforme), ☐ NOK (anomalie) ou ☐ N/A (non applicable sur cet environnement). '
        'Un chapitre n’est validé que si tous les contrôles obligatoires sont OK.',
    )
    add_para(doc, 'Conventions', bold=True)
    make_table(
        doc,
        ['Symbole', 'Signification'],
        [
            ['☐ OK', 'Comportement conforme à l’attendu'],
            ['☐ NOK', 'Anomalie — noter le n° de contrôle et une observation'],
            ['☐ N/A', 'Non testable ici (ex. biométrie sans HTTPS, hors RDC)'],
            ['Attendu', 'Résultat de référence, issu du comportement réel de l’application'],
        ],
        col_widths=[4.0, 13.3],
    )
    add_para(doc, 'Prérequis', bold=True)
    for t in [
        'Année scolaire nationale active (secondaire et/ou primaire).',
        'Au moins une école secondaire avec option Latin-Philosophie et une classe 1ère humanités.',
        'Un compte de chaque rôle (7 rôles) ; mots de passe connus du testeur.',
        'Navigateur à jour ; pour la biométrie : HTTPS ou http://localhost.',
        'Ne pas utiliser le même compte sur deux appareils pendant les tests de session unique.',
    ]:
        add_bullet(doc, t)
    add_callout(
        doc,
        'Ordre recommandé',
        'D’abord les tests transversaux (connexion), puis chaque rôle du plus large (Administrateur) '
        'au plus restreint (Enseignant), enfin les tests d’interdiction.',
    )

    # 2
    doc.add_heading('2. Synthèse des rôles', level=1)
    make_table(
        doc,
        ['Rôle', 'Périmètre', 'Écriture principale', 'Interdit'],
        [
            ['Administrateur', 'National', 'Référentiels, écoles, élèves, comptes, monitoring', 'Saisie des notes'],
            ['Agent National', 'National', 'Écoles, élèves, structure, documents', 'Comptes, monitoring, notes'],
            ['Agent PA', 'Province administrative', 'Consultation', 'Écriture métier, évaluation'],
            ['Agent PE', 'Province éducationnelle', 'Consultation', 'Écriture métier, évaluation'],
            ['Agent Antenne', 'Antenne', 'Consultation', 'Écriture métier, évaluation'],
            ['Administratif école', 'Son école', 'Fiche, élèves, programme, comptes école', 'Notes, catalogue national'],
            ['Enseignant', 'Sa classe titulaire', 'Notes, photo élève, bulletins', 'Création élève / matière / école'],
        ],
        col_widths=[3.4, 3.6, 5.2, 5.1],
    )

    # 3
    doc.add_heading('3. Tests transversaux (tous rôles)', level=1)
    add_para(
        doc,
        'À exécuter une première fois avec un compte Administrateur, puis à rejouer '
        'sur au moins un compte école et un compte enseignant pour la session unique et l’onglet unique.',
    )
    add_checklist(doc, TRANSVERSAL)

    # 4 Admin
    add_role_header(doc, '4', 'Administrateur', ADMIN['intro'])
    doc.add_heading('4.1 Navigation', level=2)
    add_checklist(doc, ADMIN['menus'])
    doc.add_heading('4.2 Tableau de bord', level=2)
    add_checklist(doc, ADMIN['dashboard'])
    doc.add_heading('4.3 Écoles et structure', level=2)
    add_checklist(doc, ADMIN['ecoles'])
    doc.add_heading('4.4 Élèves', level=2)
    add_checklist(doc, ADMIN['eleves'])
    doc.add_heading('4.5 Comptes et permissions', level=2)
    add_checklist(doc, ADMIN['users'])
    doc.add_heading('4.6 Monitoring', level=2)
    add_checklist(doc, ADMIN['monitoring'])
    doc.add_heading('4.7 Paramètres, documents et rapports', level=2)
    add_checklist(doc, ADMIN['params'])

    # 5
    add_role_header(doc, '5', 'Agent National', AGENT_NAT['intro'])
    add_checklist(doc, AGENT_NAT['items'])

    # 6
    add_role_header(doc, '6', 'Agent Province administrative', AGENT_PA['intro'])
    add_checklist(doc, AGENT_PA['items'])

    # 7
    add_role_header(doc, '7', 'Agent Province éducationnelle', AGENT_PE['intro'])
    add_checklist(doc, AGENT_PE['items'])

    # 8
    add_role_header(doc, '8', 'Agent Antenne', AGENT_ANTENNE['intro'])
    add_checklist(doc, AGENT_ANTENNE['items'])

    # 9
    add_role_header(doc, '9', 'Administratif école', ADMIN_ECOLE['intro'])
    doc.add_heading('9.1 Navigation et fiche école', level=2)
    add_checklist(doc, ADMIN_ECOLE['menus'] + ADMIN_ECOLE['ecole'])
    doc.add_heading('9.2 Élèves de l’établissement', level=2)
    add_checklist(doc, ADMIN_ECOLE['eleves'])
    doc.add_heading('9.3 Évaluation (sans saisie des notes)', level=2)
    add_checklist(doc, ADMIN_ECOLE['eval'])
    doc.add_heading('9.4 Rapports', level=2)
    add_checklist(doc, ADMIN_ECOLE['rapports'])

    # 10
    add_role_header(doc, '10', 'Enseignant', ENSEIGNANT['intro'])
    add_checklist(doc, ENSEIGNANT['items'])

    # 11
    doc.add_heading('11. Tests d’interdiction et non-régression', level=1)
    add_para(
        doc,
        'Ces contrôles vérifient qu’un utilisateur ne peut pas contourner l’interface '
        '(URL directe, second onglet, second appareil).',
    )
    add_checklist(doc, INTERDITS)
    add_callout(
        doc,
        'Latin-Philosophie',
        'Point de vigilance : pour « Littéraire · Latin-Philosophie — 1ère humanités », '
        'le menu Matière (cours de la classe) ne doit jamais rester vide ni afficher les cours '
        'd’une autre classe (ex. COURS TEST) ou le catalogue Coupe et Couture.',
    )

    # 12
    doc.add_heading('12. Fiche de recette et visas', level=1)
    add_para(doc, 'Bilan global (à remplir en fin de campagne)', bold=True)
    make_table(
        doc,
        ['Chapitre', 'OK', 'NOK', 'N/A', 'Commentaire'],
        [
            ['3. Transversal', '', '', '', ''],
            ['4. Administrateur', '', '', '', ''],
            ['5. Agent National', '', '', '', ''],
            ['6. Agent PA', '', '', '', ''],
            ['7. Agent PE', '', '', '', ''],
            ['8. Agent Antenne', '', '', '', ''],
            ['9. Administratif école', '', '', '', ''],
            ['10. Enseignant', '', '', '', ''],
            ['11. Interdictions', '', '', '', ''],
        ],
        col_widths=[5.5, 1.8, 1.8, 1.8, 6.4],
    )
    add_para(
        doc,
        'Verdict de recette :  ☐  Accepté     ☐  Accepté avec réserves     ☐  Refusé',
        bold=True,
    )
    add_para(doc, 'Liste des anomalies (n° de contrôle, rôle, description, criticité H/M/B) :')
    make_table(
        doc,
        ['N°', 'Rôle', 'Description', 'Criticité', 'Statut'],
        [['', '', '', '', '']] * 6,
        col_widths=[1.5, 3.2, 7.5, 2.3, 2.8],
    )
    add_para(doc, 'Approbation', bold=True)
    make_table(
        doc,
        ['Rôle', 'Nom', 'Date', 'Visa'],
        [
            ['Testeur', '', '', ''],
            ['Responsable métier', '', '', ''],
            ['Responsable technique', '', '', ''],
            ['Maîtrise d’ouvrage', '', '', ''],
        ],
        col_widths=[4.5, 5.0, 3.5, 4.3],
    )
    add_para(
        doc,
        f'Document généré automatiquement à partir du comportement réel d’Educ_RDC '
        f'(manuel de test v{VERSION} — {DATE_DOC}).',
        size=9,
        color=GRIS,
    )

    out_dir = Path(__file__).resolve().parent
    out = out_dir / 'Educ_RDC_Manuel_de_Test.docx'
    try:
        doc.save(out)
    except PermissionError:
        out = out_dir / f'Educ_RDC_Manuel_de_Test_v{VERSION}.docx'
        doc.save(out)
        print('Fichier principal verrouillé (ouvert dans Word ?) — sauvegarde alternative.')
    print(f'Document généré : {out}')
    return out


if __name__ == '__main__':
    build()
