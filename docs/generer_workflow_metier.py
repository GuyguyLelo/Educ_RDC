"""
Génère le document Word professionnel du workflow métier Educ_RDC.
Usage: python docs/generer_workflow_metier.py
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# Couleurs institutionnelles (alignées UI Educ_RDC)
BLEU_NUIT = RGBColor(0x0B, 0x2A, 0x4A)
BLEU = RGBColor(0x1A, 0x56, 0x8C)
GRIS = RGBColor(0x5A, 0x67, 0x72)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
VERT = RGBColor(0x1F, 0x7A, 0x4D)

VERSION = '1.2'
DATE_DOC = date.today().strftime('%d/%m/%Y')


def set_run_font(run, size=11, bold=False, color=None, name='Calibri'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn('w:shd'):
            tcPr.remove(child)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_cell_border(cell, color='1A568C', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=GRIS)


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size, color in (
        (1, 16, BLEU_NUIT),
        (2, 13, BLEU),
        (3, 11.5, BLEU_NUIT),
    ):
        style = doc.styles[f'Heading {level}']
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(8)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f'Educ_RDC  ·  Workflow métier v{VERSION}  ·  Confidentialité interne')
    set_run_font(run, size=9, color=GRIS)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run('République Démocratique du Congo  —  Identification scolaire nationale  —  Page ')
    set_run_font(r1, size=9, color=GRIS)
    add_page_number(fp)


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A568C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()

    bandeau = doc.add_paragraph()
    bandeau.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bandeau.add_run('RÉPUBLIQUE DÉMOCRATIQUE DU CONGO')
    set_run_font(r, size=12, bold=True, color=BLEU_NUIT)

    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous.add_run('Système national d’identification scolaire')
    set_run_font(r, size=11, color=BLEU)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_line(line)

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.paragraph_format.space_before = Pt(28)
    r = titre.add_run('PROCESUS MÉTIER COMPLET')
    set_run_font(r, size=26, bold=True, color=BLEU_NUIT, name='Calibri Light')

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous_titre.add_run('Workflow opérationnel de la plateforme Educ_RDC')
    set_run_font(r, size=14, color=BLEU)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(36)
    r = meta.add_run(
        f'Document de référence\nVersion {VERSION}  ·  {DATE_DOC}\n'
        'Destination : maîtrise d’ouvrage, exploitation et formation'
    )
    set_run_font(r, size=11, color=GRIS)

    for _ in range(4):
        doc.add_paragraph()

    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = box.add_run(
        'Ce document décrit les acteurs, la hiérarchie territoriale,\n'
        'les processus de bout en bout et les règles métier applicables.'
    )
    set_run_font(r, size=10, color=GRIS)

    doc.add_page_break()


def add_toc_like(doc: Document):
    doc.add_heading('Sommaire', level=1)
    items = [
        '1. Objet et périmètre de la plateforme',
        '2. Acteurs et responsabilités',
        '3. Architecture organisationnelle',
        '4. Vue d’ensemble du workflow',
        '5. Processus détaillés',
        '   5.1 Paramétrage du référentiel territorial',
        '   5.2 Identification des écoles',
        '   5.3 Structure scolaire et matières',
        '   5.4 Enregistrement des élèves et des parents',
        '   5.5 Identification du personnel scolaire',
        '   5.6 Gestion des comptes utilisateurs',
        '   5.7 Module Évaluation (école)',
        '   5.8 Biométrie',
        '   5.9 Cartes scolaires (API / fiche élève)',
        '   5.10 Imports Excel',
        '   5.11 Pilotage, tableaux de bord et rapports',
        '   5.12 Gestion de permissions (admin)',
        '6. Matrice RACI synthétique',
        '7. Règles métier structurantes',
        '8. Modules applicatifs et navigation',
        '9. Parcours type de mise en service',
        '10. Glossaire',
        '11. Historique des versions',
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=11, color=BLEU_NUIT)
    doc.add_page_break()


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10, bold=True, color=BLANC)
        shade_cell(cell, '0B2A4A')
        set_cell_border(cell, '0B2A4A', '6')
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_run_font(r, size=9.5)
            if r_idx % 2 == 1:
                shade_cell(cell, 'F4F7FA')
            set_cell_border(cell, 'C5D0DC', '4')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    r = p.add_run(f'▶ {title}  ')
    set_run_font(r, size=10, bold=True, color=BLEU)
    r2 = p.add_run(text)
    set_run_font(r2, size=10, color=GRIS)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)


def build():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_toc_like(doc)

    # 1
    doc.add_heading('1. Objet et périmètre de la plateforme', level=1)
    add_para(
        doc,
        'Educ_RDC est la plateforme nationale d’identification scolaire de la '
        'République Démocratique du Congo. Elle assure le recensement et le suivi '
        'des établissements, des classes, des élèves (identité, parents, photo), '
        'du personnel scolaire, de la biométrie, des cartes scolaires et de '
        'l’évaluation scolaire (notes / bulletins), avec un pilotage territorial par rôle.'
    )
    add_para(doc, 'Objectifs opérationnels', bold=True)
    for t in [
        'Structurer le référentiel territorial (PA → PE → Antenne → École → Classe).',
        'Identifier de façon fiable les écoles et leurs effectifs.',
        'Enregistrer les élèves avec identité complète et contacts parentaux.',
        'Piloter l’évaluation scolaire au niveau de l’école (titulaire + administratif).',
        'Gérer les cartes scolaires via l’API et la fiche élève (PDF / QR).',
        'Donner à chaque acteur une vue adaptée à son périmètre (consultation ou écriture).',
        'Accélérer le déploiement via imports Excel contrôlés.',
    ]:
        add_bullet(doc, t)
    add_callout(
        doc,
        'Navigation v1.2',
        'Le menu latéral est défilable. Le module « Cartes » n’apparaît pas dans le menu : '
        'les cartes restent accessibles depuis la fiche élève et l’API. '
        'Le module « Évaluation » est réservé à l’administratif école et à l’enseignant.',
    )

    # 2
    doc.add_heading('2. Acteurs et responsabilités', level=1)
    add_para(
        doc,
        'L’accès est authentifié. Les droits dépendent du rôle métier rattaché '
        'au compte utilisateur. Les agents territoriaux (PA, PE, antenne) partagent '
        'les mêmes restrictions d’écriture (consultation seule) via le marqueur '
        'technique est_agent_territorial.'
    )
    make_table(
        doc,
        ['Rôle', 'Périmètre', 'Droits principaux', 'Accès UI notable'],
        [
            [
                'Administrateur (admin)',
                'National',
                'Pilotage + comptes + monitoring + catalogue matières ; pas d’Évaluation',
                'Utilisateurs, Permissions, Paramètres, Gestion documentaire',
            ],
            [
                'Agent national',
                'National',
                'Écriture écoles/élèves/référentiels ; pas de comptes ni Évaluation',
                'Paramètres, documentaire (pas Gestion Utilisateurs)',
            ],
            [
                'Agent Province administrative',
                'PA rattachée',
                'Consultation seule — vue d’ensemble des PE',
                'Dashboard PE, Écoles/Élèves/Rapports (lecture)',
            ],
            [
                'Agent Province éducationnelle',
                'PE rattachée',
                'Consultation seule — vue d’ensemble des antennes',
                'Dashboard antennes, Écoles/Élèves/Rapports (lecture)',
            ],
            [
                'Agent antenne',
                'Antenne',
                'Consultation seule — effectifs par école',
                'Dashboard écoles (paginé), Écoles/Élèves/Rapports (lecture)',
            ],
            [
                'Administratif école (admin_ecole)',
                'Son école',
                'Fiche école, élèves, programme, comptes école ; pas de saisie notes ni création matières',
                'Mon école, Élèves, Évaluation (config), Rapports',
            ],
            [
                'Enseignant',
                'Classe titulaire',
                'Lecture élèves + photo + saisie notes + bulletins / listes PDF',
                'Élèves, Évaluation (saisie), Rapports classe',
            ],
        ],
        col_widths=[3.4, 2.8, 5.0, 4.3],
    )
    add_callout(
        doc,
        'Hiérarchie des agents',
        'National → Agent Province administrative → Agent Province éducationnelle → '
        'Agent Antenne → Administratif école / Enseignant. '
        'L’agent PA est le chef hiérarchique des agents PE.',
    )
    add_callout(
        doc,
        'Règle d’écriture territoriale',
        'Les agents PA / PE / antenne ne créent ni ne modifient écoles, élèves, photos, '
        'documents, personnel, ni n’accèdent à l’Évaluation. '
        'La saisie des notes est réservée à l’enseignant titulaire.',
    )

    # 3
    doc.add_heading('3. Architecture organisationnelle', level=1)
    add_para(
        doc,
        'La hiérarchie territoriale constitue le socle du filtrage des données '
        'et des workflows. Les libellés de référence peuvent être préfixés '
        '(ex. « Prov. Admin - … », « Prov. Educ. - … »).'
    )
    make_table(
        doc,
        ['Niveau', 'Entité', 'Rattachement', 'Rôle clé'],
        [
            ['1', 'Province administrative (PA)', '—', 'Agent province administrative'],
            ['2', 'Province éducationnelle (PE)', 'PA', 'Agent province éducationnelle'],
            ['3', 'Antenne', 'PE', 'Agent antenne'],
            ['4', 'École', 'PE + Antenne', 'Administratif école'],
            ['5', 'Classe', 'École', 'Enseignant titulaire'],
            ['6', 'Élève', 'École + Classe', 'Dossier scolaire'],
        ],
        col_widths=[1.5, 4.5, 3.5, 5.5],
    )
    add_para(doc, 'Rattachements des utilisateurs', bold=True)
    for t in [
        'Agent province admin. : rattachement PA uniquement (PE / antenne / école vidés).',
        'Agent provincial : rattachement PE.',
        'Agent antenne : rattachement antenne.',
        'Administratif école et enseignant : rattachement à une école.',
        'Enseignant : rattachement obligatoire à une classe titulaire de son école.',
    ]:
        add_bullet(doc, t)

    # 4
    doc.add_heading('4. Vue d’ensemble du workflow', level=1)
    add_para(
        doc,
        'Le processus global suit un ordre de dépendance strict. Le non-respect '
        'de cet ordre bloque les étapes aval (ex. : élève sans classe existante).'
    )
    make_table(
        doc,
        ['Étape', 'Processus', 'Prérequis', 'Résultat'],
        [
            ['1', 'Paramétrage territorial', 'Compte national', 'PA, PE, Antennes actives'],
            ['2', 'Identification des écoles', 'Antenne / PE existants', 'Écoles référencées'],
            ['3', 'Structure scolaire & matières', 'École + catalogue national', 'Sections / options / classes / matières'],
            ['4', 'Enregistrement élèves & parents', 'Classe existante', 'Dossiers élèves complets'],
            ['5', 'Personnel scolaire', 'École existante', 'Effectifs RH identifiés'],
            ['6', 'Comptes utilisateurs école', 'École (+ classe pour enseignant)', 'Accès applicatifs'],
            ['7', 'Évaluation (programme + notes)', 'Matières + classe + titulaire', 'Notes, bulletins PDF'],
            ['8', 'Biométrie & cartes', 'Élève (+ photo)', 'Dossier biométrique / carte'],
            ['9', 'Pilotage & rapports', 'Données opérationnelles', 'Indicateurs & PDF'],
        ],
        col_widths=[1.5, 4.2, 4.3, 5.0],
    )

    # 5
    doc.add_heading('5. Processus détaillés', level=1)

    doc.add_heading('5.1 Paramétrage du référentiel territorial', level=2)
    add_para(doc, 'Objectif : constituer le référentiel national PA / PE / Antennes.', bold=True)
    add_para(doc, 'Acteurs : Administrateur, Agent national.')
    add_para(doc, 'Enchaînement', bold=True)
    for t in [
        'Se connecter avec un compte national, ouvrir le menu Paramètres.',
        'Créer les provinces administratives (nom, code unique).',
        'Créer les provinces éducationnelles rattachées à une PA.',
        'Créer les antennes rattachées à une PE (adresse, téléphone).',
        'Contrôler la cohérence via l’organigramme.',
        'Modifier ou supprimer depuis la modale de détail (jamais depuis la liste).',
    ]:
        add_numbered(doc, t)
    add_callout(
        doc,
        'Ordre imposé',
        'PA → PE → Antennes → Écoles. Une antenne ne peut exister sans PE ; '
        'une école doit être rattachée à une PE et une antenne.',
    )

    doc.add_heading('5.2 Identification des écoles', level=2)
    add_para(doc, 'Objectif : référencer chaque établissement scolaire.', bold=True)
    add_para(
        doc,
        'Acteurs écriture : Admin, Agent national, Administratif école (sa fiche). '
        'Agents territoriaux : consultation seule.',
    )
    add_para(doc, 'Enchaînement', bold=True)
    for t in [
        'Accéder à la liste Écoles (filtrée selon le périmètre territorial).',
        'Créer une école (national) : identité, code, type, niveau, adresse, contacts, GPS, rattachement PE/Antenne.',
        'Compléter la fiche : photos, documents, localisation, effectifs, personnel.',
        'Pour l’administratif école : accès exclusif via « Mon école ».',
        'Agents PA / PE / antenne : consultation des fiches sans boutons d’écriture.',
    ]:
        add_numbered(doc, t)
    add_para(doc, 'Règles', bold=True)
    for t in [
        'L’enseignant n’a aucun accès aux écrans Écoles.',
        'L’administratif école ne voit et ne gère que son établissement.',
        'Les agents territoriaux ne voient que les écoles de leur PA / PE / antenne.',
        'Les agents territoriaux ne peuvent ni créer, ni modifier, ni supprimer une école.',
    ]:
        add_bullet(doc, t)

    doc.add_heading('5.3 Structure scolaire et matières', level=2)
    add_para(
        doc,
        'Objectif : définir sections, options, classes et catalogue de matières.',
        bold=True,
    )
    add_para(doc, 'Acteurs', bold=True)
    for t in [
        'Structure (sections / options / classes) : administration nationale (Paramètres → Structure scolaire) '
        'et gestion opérationnelle école selon les droits existants.',
        'Catalogue des matières : création / catalogue réservés à l’administration nationale.',
        'Administratif école : applique le programme de classe à partir des matières déjà définies '
        '(sans créer de nouvelles matières).',
    ]:
        add_bullet(doc, t)
    add_callout(
        doc,
        'Point de contrôle',
        'Sans matières configurées au niveau national pour la section / option, '
        'l’administratif école ne peut pas constituer le programme de classe.',
    )

    doc.add_heading('5.4 Enregistrement des élèves et des parents', level=2)
    add_para(doc, 'Objectif : constituer le dossier scolaire complet de l’élève.', bold=True)
    add_para(
        doc,
        'Acteurs écriture : Admin, Agent national, Administratif école. '
        'Agents territoriaux : consultation. Enseignant : consultation + photo de sa classe.',
    )
    add_para(doc, 'Enchaînement', bold=True)
    for t in [
        'Accéder à Élèves (ou depuis la fiche école).',
        'Saisir l’identité : matricule, N° identification / permanent, nom, postnom, prénom, sexe, naissance.',
        'Rattacher l’élève à une école et à une classe existante.',
        'Renseigner l’adresse et les parents (père, mère, tuteur + photos / contacts).',
        'Déposer la photo de l’élève (initialise la biométrie).',
        'Importer en masse via Excel (national / admin école) après création des classes.',
    ]:
        add_numbered(doc, t)
    make_table(
        doc,
        ['Bloc', 'Informations gérées'],
        [
            ['Identité élève', 'Matricule, N° ID / permanent, nom, postnom, prénom, sexe, date et lieu de naissance, photo'],
            ['Scolarité', 'École, classe obligatoire, adresse'],
            ['Père / Mère / Tuteur', 'Identité, téléphone, e-mail, profession, photo ; lien de parenté pour le tuteur'],
        ],
        col_widths=[3.5, 12.0],
    )

    doc.add_heading('5.5 Identification du personnel scolaire', level=2)
    add_para(
        doc,
        'Objectif : identifier le personnel RH de l’école (distinct des comptes de connexion).',
        bold=True,
    )
    add_para(doc, 'Écriture : national et administratif école. Agents territoriaux : lecture seule.')
    for t in [
        'Ouvrir la fiche école → section Personnel.',
        'Identifier un agent : matricule, identité, fonction, contacts.',
        'Importer un fichier Excel/CSV si besoin (hors agent territorial).',
    ]:
        add_numbered(doc, t)

    doc.add_heading('5.6 Gestion des comptes utilisateurs', level=2)
    add_para(doc, 'Objectif : délivrer les accès applicatifs selon le rôle et le périmètre.', bold=True)
    add_para(doc, 'Acteur exclusif pour les comptes agents / admin : Administrateur national.')
    add_para(doc, 'Administratif école : peut gérer les comptes admin_ecole / enseignant de son école (fiche école).')
    add_para(doc, 'Enchaînement', bold=True)
    for t in [
        'Ouvrir Gestion Utilisateurs (admin).',
        'Créer un agent PA : rôle agent_province_admin + province administrative.',
        'Créer un agent PE : rôle agent_provincial + province éducationnelle.',
        'Créer un agent antenne : rôle agent_antenne + antenne.',
        'Créer un administratif école : admin_ecole + école.',
        'Créer un enseignant : école + classe titulaire (via fiche personnel enseignant).',
        'Activer / désactiver ; réinitialiser le mot de passe si besoin.',
    ]:
        add_numbered(doc, t)

    doc.add_heading('5.7 Module Évaluation (école)', level=2)
    add_para(
        doc,
        'Objectif : saisir les notes, constituer le programme de classe et éditer les bulletins.',
        bold=True,
    )
    add_para(doc, 'Accès page / API : uniquement administratif école et enseignant. '
             'National et agents territoriaux : refusés.')
    make_table(
        doc,
        ['Action', 'Administratif école', 'Enseignant'],
        [
            ['Voir le module Évaluation', 'Oui', 'Oui (classe figée)'],
            ['Appliquer le programme de classe', 'Oui', 'Non'],
            ['Créer / charger le catalogue matières', 'Non (national)', 'Non'],
            ['Saisir / enregistrer les notes', 'Non', 'Oui (titulaire)'],
            ['Déverrouiller une période', 'Oui', 'Non'],
            ['Actualiser le classement / décisions', 'Oui', 'Oui (sa classe)'],
            ['Imprimer bulletins / listes PDF', 'Oui (école)', 'Oui (classe)'],
        ],
        col_widths=[5.5, 5.0, 5.0],
    )
    add_callout(
        doc,
        'Saisie des notes',
        'La grille reste consultable par l’administratif école en lecture seule. '
        'Seul le titulaire peut ouvrir une période de saisie et enregistrer les notes.',
    )

    doc.add_heading('5.8 Biométrie', level=2)
    add_para(doc, 'Objectif : constituer le dossier biométrique de l’élève.', bold=True)
    for t in [
        'La photo élève crée ou met à jour automatiquement l’enregistrement biométrique.',
        'Une empreinte de référence est générée (hash technique) et peut être validée via l’API.',
        'La fiche élève affiche le statut (validée / en attente) et la date de capture.',
        'Les agents territoriaux consultent sans modifier ; l’enseignant met à jour la photo de sa classe.',
    ]:
        add_numbered(doc, t)

    doc.add_heading('5.9 Cartes scolaires (API / fiche élève)', level=2)
    add_para(doc, 'Objectif : produire et consulter les cartes scolaires sécurisées.', bold=True)
    add_callout(
        doc,
        'UI',
        'Pas de menu « Cartes ». Consultation depuis la fiche élève (PDF / QR). '
        'Émission via l’API. Blocs cartes / biométries masqués pour admin_ecole / enseignant selon règles UI.',
    )
    for t in [
        'Émission via l’API (numéro type RDC-…, expiration +3 ans, QR).',
        'Consultation PDF / QR depuis la fiche élève (rôles autorisés).',
        'Statuts : active, expirée, annulée.',
    ]:
        add_numbered(doc, t)

    doc.add_heading('5.10 Imports Excel', level=2)
    add_para(doc, 'Objectif : accélérer le déploiement massif avec contrôle qualité.', bold=True)
    make_table(
        doc,
        ['Modèle', 'Qui importe', 'Notes'],
        [
            ['PA / PE / Antennes', 'National', 'Référentiel'],
            ['Écoles', 'National', 'Après antennes'],
            ['Classes', 'National / admin école', 'Avant les élèves'],
            ['Élèves (+ parents)', 'National / admin école', 'Agents territoriaux : non'],
            ['Personnel', 'National / admin école', 'Agents territoriaux : non'],
        ],
        col_widths=[4.0, 4.5, 7.0],
    )

    doc.add_heading('5.11 Pilotage, tableaux de bord et rapports', level=2)
    add_para(doc, 'Objectif : suivre les indicateurs selon le périmètre de l’acteur.', bold=True)
    make_table(
        doc,
        ['Rôle', 'Portée du dashboard', 'Spécificités'],
        [
            ['Admin / Agent national', 'Nationale', 'KPI élèves / écoles / cartes / biométries'],
            ['Agent province admin.', 'PA', 'KPI PE / écoles / élèves ; tableau effectifs par PE'],
            ['Agent provincial', 'PE', 'KPI antennes / écoles / élèves ; tableau par antenne'],
            ['Agent antenne', 'Antenne', 'KPI écoles / élèves ; tableau par école (pagination 20)'],
            ['Administratif école', 'École', 'Effectifs + personnel'],
            ['Enseignant', 'Classe', 'Effectif H/F de la classe'],
        ],
        col_widths=[3.8, 3.5, 8.2],
    )
    add_para(
        doc,
        'Rapports : les agents territoriaux ne voient pas les KPI cartes / biométries. '
        'L’agent PA affiche le KPI « Provinces éduc. » ; l’agent PE affiche « Antennes ». '
        'Export PDF disponible selon le périmètre.'
    )

    doc.add_heading('5.12 Gestion de permissions (admin)', level=2)
    add_para(
        doc,
        'Objectif : consulter la matrice des permissions réellement implémentées dans le code.',
        bold=True,
    )
    for t in [
        'Accessible uniquement à l’administrateur national (Utilisateurs → Gestion de permissions).',
        'Affiche la matrice capacités × rôles et le détail par rôle.',
        'Bouton « Imprimer PDF » : export paysage de la matrice + détail des rôles.',
        'Page de consultation : ne permet pas encore de modifier les droits en base.',
    ]:
        add_numbered(doc, t)

    # 6 RACI
    doc.add_heading('6. Matrice RACI synthétique', level=1)
    add_para(
        doc,
        'R = Réalise  ·  A = Approuve  ·  C = Consulté  ·  I = Informé  ·  L = Lecture seule  ·  — = Non concerné'
    )
    make_table(
        doc,
        ['Processus', 'Admin', 'Nat.', 'PA', 'PE', 'Ant.', 'Adm. éc.', 'Ens.'],
        [
            ['Paramétrage PA/PE/Antenne', 'A/R', 'R', '—', '—', '—', '—', '—'],
            ['Catalogue matières', 'A/R', 'R', '—', '—', '—', '—', '—'],
            ['Identification écoles', 'A', 'R', 'L', 'L', 'L', 'R (sienne)', '—'],
            ['Élèves & parents', 'A', 'R', 'L', 'L', 'L', 'R', 'I (+ photo)'],
            ['Personnel scolaire', 'A', 'R', 'L', 'L', 'L', 'R', '—'],
            ['Comptes agents / admin', 'A/R', '—', '—', '—', '—', '—', '—'],
            ['Comptes école / enseignant', 'A', '—', '—', '—', '—', 'R', '—'],
            ['Programme de classe', '—', '—', '—', '—', '—', 'R', 'I'],
            ['Saisie des notes', '—', '—', '—', '—', '—', 'L', 'R'],
            ['Bulletins / classement', '—', '—', '—', '—', '—', 'R', 'R'],
            ['Biométrie / cartes', 'A', 'R', 'L', 'L', 'L', 'partiel', 'partiel'],
            ['Dashboard / rapports', 'A/R', 'R', 'L', 'L', 'L', 'R', 'R'],
            ['Matrice permissions PDF', 'A/R', '—', '—', '—', '—', '—', '—'],
        ],
        col_widths=[3.6, 1.5, 1.4, 1.3, 1.3, 1.3, 2.0, 1.8],
    )

    # 7
    doc.add_heading('7. Règles métier structurantes', level=1)
    rules = [
        'Les agents PA, PE et antenne sont en consultation territoriale (est_agent_territorial) : pas d’écriture écoles/élèves/personnel/import.',
        'L’agent province administrative est le chef hiérarchique des agents PE ; son dashboard consolide les PE.',
        'L’administrateur et l’agent national n’accèdent pas au module Évaluation.',
        'L’administratif école n’ajoute pas de matières et ne saisit pas les notes.',
        'Seul l’enseignant titulaire saisit les notes de sa classe.',
        'Le catalogue matières est géré par l’administration nationale (Structure scolaire).',
        'L’enseignant ne voit que les élèves de sa classe titulaire ; sans classe, aucune donnée élève.',
        'L’enseignant n’a pas accès au menu Écoles / Mon école.',
        'L’administratif école n’accède qu’à « Mon école » et aux données de son établissement.',
        'La classe est obligatoire pour tout élève et doit appartenir à son école.',
        'La classe titulaire est obligatoire pour tout compte enseignant.',
        'Les suppressions UI s’effectuent depuis la fiche détail ou la modale, jamais depuis les listes.',
        'La photo élève initialise / synchronise le dossier biométrique.',
        'Les cartes portent un numéro automatique, une durée de validité et un QR ; pas de menu Cartes dédié.',
        'La création des comptes agents / administrateur est réservée à l’admin national.',
        'La page Gestion de permissions (avec export PDF) est réservée à l’admin national.',
        'Le menu latéral est défilable lorsque la liste des entrées dépasse la hauteur d’écran.',
    ]
    for i, rule in enumerate(rules, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'R{i:02d}. ')
        set_run_font(r, size=11, bold=True, color=BLEU)
        r2 = p.add_run(rule)
        set_run_font(r2, size=11)

    # 8
    doc.add_heading('8. Modules applicatifs et navigation', level=1)
    add_para(doc, 'Menu latéral (ordre d’affichage)', bold=True)
    make_table(
        doc,
        ['Entrée menu', 'Visibilité', 'Fonction'],
        [
            ['Tableau de bord', 'Tous les rôles', 'Pilotilotage et indicateurs'],
            ['Mon école', 'Administratif école uniquement', 'Fiche de son établissement'],
            ['Écoles', 'Tous sauf enseignant et admin_ecole', 'Liste / consultation / identification'],
            ['Élèves', 'Tous (enseignant = sa classe)', 'Dossiers, parents, photo, import'],
            ['Évaluation', 'admin_ecole + enseignant', 'Programme, notes, bulletins'],
            ['Gestion documentaire', 'Admin + Agent national', 'Référentiel documentaire'],
            ['Rapports', 'Tous les rôles', 'Analyse et exports'],
            ['Gestion Utilisateurs', 'Administrateur uniquement', 'Comptes et rattachements'],
            ['Utilisateurs connectés', 'Administrateur uniquement', 'Monitoring sessions'],
            ['Paramètres', 'Admin + Agent national', 'Territoire, structure, années'],
        ],
        col_widths=[4.0, 5.0, 6.5],
    )
    add_para(
        doc,
        'Hors menu : cartes depuis la fiche élève ; API REST / JWT pour opérations avancées. '
        'Sous Utilisateurs : page Gestion de permissions + impression PDF.'
    )

    # 9
    doc.add_heading('9. Parcours type de mise en service', level=1)
    add_para(doc, 'Scénario recommandé pour ouvrir une province / antenne / école en production :')
    steps = [
        ('J0 — Gouvernance', 'Créer les comptes admin et agents nationaux ; valider le découpage PA/PE.'),
        ('J1 — Référentiel', 'Saisir PA, PE, Antennes ; contrôler l’organigramme.'),
        ('J1b — Agents territoriaux', 'Créer agent PA, agents PE, agents antenne (consultation).'),
        ('J2 — Établissements', 'Identifier les écoles (saisie ou import national).'),
        ('J3 — Structure & matières', 'Sections / options / classes ; catalogue matières national.'),
        ('J4 — Élèves', 'Importer ou saisir les élèves avec parents et photos.'),
        ('J5 — RH & accès école', 'Personnel ; comptes admin_ecole et enseignants.'),
        ('J6 — Évaluation', 'Admin école applique le programme ; titulaire saisit les notes.'),
        ('J7 — Biométrie & cartes', 'Photos, biométries ; émettre les cartes (API).'),
        ('J8 — Pilotage', 'Former chaque rôle sur son dashboard et ses restrictions.'),
    ]
    for title, text in steps:
        p = doc.add_paragraph()
        r = p.add_run(f'{title} — ')
        set_run_font(r, size=11, bold=True, color=VERT)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)

    # 10
    doc.add_heading('10. Glossaire', level=1)
    make_table(
        doc,
        ['Terme', 'Définition'],
        [
            ['PA', 'Province administrative'],
            ['PE', 'Province éducationnelle'],
            ['Agent territorial', 'Agent PA, PE ou antenne — consultation seule'],
            ['Agent province admin.', 'Chef hiérarchique des agents PE (rôle agent_province_admin)'],
            ['Administratif école', 'Compte de gestion de l’établissement (admin_ecole)'],
            ['Classe titulaire', 'Classe dont l’enseignant est responsable dans l’application'],
            ['Catalogue matières', 'Référentiel de branches géré par le national'],
            ['Programme de classe', 'Attachement des matières à une classe / année'],
            ['Matricule', 'Identifiant unique de l’élève'],
            ['Biométrie', 'Dossier photo / empreinte associé à l’élève'],
            ['Carte scolaire', 'Support d’identification avec QR ; fiche élève / API'],
            ['Matrice permissions', 'Inventaire des droits réellement implémentés (page + PDF)'],
            ['Périmètre', 'Sous-ensemble de données visible selon le rôle'],
        ],
        col_widths=[4.0, 11.5],
    )

    # 11 Historique
    doc.add_heading('11. Historique des versions', level=1)
    make_table(
        doc,
        ['Version', 'Date', 'Évolutions'],
        [
            ['1.0', '05/08/2026', 'Version initiale du workflow métier Educ_RDC'],
            [
                '1.1',
                '14/08/2026',
                'Retrait du menu Cartes ; menu latéral fixe ; '
                'enseignant sans accès Mon école ; parents enrichis ; '
                'suppressions depuis détail/modales',
            ],
            [
                '1.2',
                DATE_DOC,
                'Agent Province administrative ; agents territoriaux en lecture seule ; '
                'Évaluation réservée à l’école ; admin école sans saisie notes ni création matières ; '
                'catalogue matières national ; page Gestion de permissions + PDF ; '
                'dashboards PA/PE/antenne ; menu latéral défilable',
            ],
        ],
        col_widths=[2.2, 2.8, 10.5],
    )

    doc.add_heading('Approbation', level=1)
    add_para(
        doc,
        f'Document généré à partir du fonctionnement réel de la plateforme Educ_RDC (v{VERSION}). '
        'Il constitue la référence métier pour la formation, l’exploitation et la recette.'
    )
    make_table(
        doc,
        ['Rôle', 'Nom', 'Date', 'Visa'],
        [
            ['Maîtrise d’ouvrage', '', '', ''],
            ['Responsable métier', '', '', ''],
            ['Responsable technique', '', '', ''],
        ],
        col_widths=[4.5, 4.5, 3.0, 3.5],
    )

    out_dir = Path(__file__).resolve().parent
    out = out_dir / 'Educ_RDC_Workflow_Metier.docx'
    try:
        doc.save(out)
    except PermissionError:
        out = out_dir / f'Educ_RDC_Workflow_Metier_v{VERSION}.docx'
        doc.save(out)
        print('Fichier principal verrouillé (ouvert dans Word ?) — sauvegarde alternative.')
    print(f'Document généré : {out}')
    return out


if __name__ == '__main__':
    build()
